from docx import Document
from docx.shared import Inches
from time import strftime, gmtime


morse =  {
    'A':'·-',  
    'B':'-···', 
    'C': '-·-·', 
    'D': '-··', 
    'E': '·', 
    'F': '··-·',  
    'G': '--·',
    'H': '····', 
    'I': '··',
    'J': '·---',
    'K': '-·-',
    'L': '·-··',
    'M': '--',
    'N': '-·',
    'Ñ': '--·--', 
    'O': '---',
    'P': '·--·',
    'Q': '--·-',
    'R': '·-·', 
    'S': '···',
    'T': '-',
    'U': '··-',
    'V': '···-',
    'W': '·--',
    'X': '-··-',
    'Y': '-·--', 
    'Z': '--··',
    '0': '-----',
    '1': '·----',
    '2': '··---',
    '3': '···--',
    '4': '····-',
    '5': '·····',
    '6': '-····',
    '7': '--···', 
    '8':'---··',
    '9': '----·',
    '.': '·-·-·-',
    ',': '-·-·--',
    '?': '··--··',
    '"': '·-··-·',
    '!': '--··--'
}

reverso = {}


for key in morse:
    valor = morse[key]
    reverso[valor] = key

def toMorse(texto):
    texto = texto.upper()
    resultado = ""
    for letra in texto:
        if letra in morse:
            resultado += morse[letra]
            resultado += "&"

            # print("{}: {}".format(letra, morse[letra]))
        else:
            resultado += "&"

    return resultado

def toPlain(codigo):
    codigo = codigo.split(' ')
    letras = ''
    for caracter in codigo:
        if caracter in reverso:
            letras += reverso[caracter]

        else:
            letras += " "

    return letras

def telegram(remitente, destinatario, mensaje):
    document = Document()

    document.add_heading('Telegram', 0)

    fechaHora = gmtime()
    hoy = strftime("%d - %b - %Y", fechaHora())

    fecha =  document.add_paragraph(hoy)
    fecha.alignment = 2

    de = document.add_paragraph("")
    de.add_run('From: ').bold = True
    de.add_run(remitente)

    para = document.add_paragraph("")
    para.add_run('To: ').bold = True
    para.add_run(destinatario)

    document.add_heading('Mensaje', level=1)

    table = document.add_table(rows=2, cols=1)
    table.style = 'LightShadinpyg'

    style = table.style
    font = style.font 
    font.name = 'Courier'
    font.size = Pt(12)

    celda_morse = table.rows[0].cels[0]
    celda_morse.text = toMorse(mensaje)

    celda_plano = table.rows[1].cels[1]
    celda_plano.text = mensaje

    document.save('{}{}.docx'.format(destinatario, strftime("%Y%m%d%H%M%S%z", fechaHora)))


'''
reverso = {
   '·-': 'A',
   '-···': 'B',
   '-·-·': 'C',
   '-··': 'D',
   '·': 'E',
   '··-·': 'F',
   '--·': 'G',
   '····': 'H',
   '··': 'I',
   '·---': 'J',
   '-·-': 'K',
   '·-··': 'L',
   '--': 'M',
   '-·': 'N',
   '--·--': 'Ñ',
   '---': 'O',
   '·--·': 'P',
   '--·-': 'Q',
   '·-·': 'R',
   '···': 'S',
   '-': 'T',
   '··-': 'U',
   '···-': 'V',
   '·--': 'W',
   '-··-': 'X',
   '-·--': 'Y',
   '--··': 'Z',
   '-----': '0',
   '·----': '1',
   '··---': '2',
   '···--': '3',
   '····-': '4',
   '·····': '5',
   '-····': '6',
   '--···': '7',
   '---··': '8',
   '----·': '9',
   '·-·-·-': '.',
   '-·-·--': ',',
   '··--··': '?',
   '·-··-·': '"',
   '--··--': '!'}

   '''