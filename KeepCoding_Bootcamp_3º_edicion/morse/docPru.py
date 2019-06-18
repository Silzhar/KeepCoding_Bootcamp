import docx
from docx import Document
from docx.shared import Inches
from time import strftime, gmtime

document = Document()

document.add_heading('Telegram', 0)

hoy = strftime("%d - %b - %Y", gmtime())

fecha =  document.add_paragraph(hoy)
fecha.alignment = 1

de = document.add_paragraph("")
de.add_run('From:').bold = True
de.add_run('Remitente')

para = document.add_paragraph("")
para.add_run('From:').bold = True
para.add_run('Destinatario')

document.add_heading('Mensaje', level=1)

table = document.add_table(rows=2, cols=1)
table.style = 'LightShadinpyg'

style = table.style
font = style.font 
font.name = 'Courier'
font.size = Pt(12)

celda_morse = table.rows[0].cels[0]
celda_morse.text = ".-.-."

celda_plano = table.rows[1].cels[1]
celda_plano.text = "Hola"

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))


document.add_page_break()

document.save('docPru')